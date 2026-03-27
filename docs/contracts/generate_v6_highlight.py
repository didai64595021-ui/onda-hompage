#!/usr/bin/env python3
"""v6 계약서 docx - 사업자번호 수정 + 클라이언트별 변동부분 노란 하이라이트"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

LIGHT_YELLOW = 'FFFF99'

def set_highlight_color(run, color_hex):
    """밝은 노란색 등 커스텀 하이라이트를 XML shading으로 적용"""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    # 기존 shd 제거
    for existing in rPr.findall(qn('w:shd')):
        rPr.remove(existing)
    # 기존 highlight 제거
    for existing in rPr.findall(qn('w:highlight')):
        rPr.remove(existing)
    rPr.append(shd)

doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# === 클라이언트별 변동 데이터 (하이라이트 대상) ===
# 이 값들만 바꾸면 다른 클라이언트 계약서 생성 가능
CLIENT = {
    'company': '마드모아젤헤어',
    'representative': '홍길동',
    'biz_number': '000-00-00000',
    'address': '세종특별자치시 OO로 OO, O층',
    'phone': '010-0000-0000',
    'email': 'sample@email.com',
    'contract_date_text': '2026년 03월 27일',
    'contract_date': '2026-03-27',
    'period_start': '2026-04-01',
    'period_end': '2026-04-30',
    'period_months': '1',
    'branches': [
        ('1', '마드모아젤헤어 마크원애비뉴점', '세종 나성동 미용실', '13위'),
        ('2', '마드모아젤헤어 나성1호점', '나성동미용실', '15위'),
        ('3', '마드모아젤헤어 나성2호점', '세종미용실', '20위'),
        ('4', '마드모아젤헤어 반곡점', '반곡동미용실', '21위'),
    ],
    'price_per_branch': '900,000',
    'price_total': '3,600,000',
    'price_3branch': '1,000,000',
    'price_2branch': '1,100,000',
    'price_1branch': '1,200,000',
}

OUR_BIZ_NUMBER = '739-40-01011'


def add_title(text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True if size > 12 else False


def add_article(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True


def add_text(text, highlight=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    if highlight:
        set_highlight_color(run, LIGHT_YELLOW)


def add_table(headers, rows, highlight_cells=None):
    """highlight_cells: set of (row_idx, col_idx) to highlight yellow"""
    if highlight_cells is None:
        highlight_cells = set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if (ri, ci) in highlight_cells:
                        set_highlight_color(r, LIGHT_YELLOW)


# ===================== 본문 시작 =====================

add_title('온라인 마케팅 계약서', 18)
add_title('(안전보장형)', 14)
doc.add_paragraph()

# 인트로 - 날짜 하이라이트
p = doc.add_paragraph()
r1 = p.add_run('"갑"과 "을"은 ')
r1.font.size = Pt(10)
r2 = p.add_run(f'[ {CLIENT["contract_date"]} ]')
r2.font.size = Pt(10)
set_highlight_color(r2, LIGHT_YELLOW)
r3 = p.add_run(' 아래와 같이 온라인 마케팅 계약을 체결한다.')
r3.font.size = Pt(10)

# === 제1조 ===
add_article('제1조 (계약의 당사자)')
add_text('① 계약의 당사자는 다음과 같다.')

# 당사자 테이블 - 을 정보 전부 하이라이트
hl_cells = set()
for ri in range(6, 12):  # 을 행 (index 6~11)
    hl_cells.add((ri, 2))
add_table(['구분', '항목', '내용'], [
    ['갑 (마케팅 제공업체)', '상호', '온다마케팅'],
    ['', '대표자', '황승현'],
    ['', '사업자등록번호', OUR_BIZ_NUMBER],
    ['', '주소', '서울시 강남구 테헤란로 82길 15, 3층 304호'],
    ['', '연락처', '010-2420-8657'],
    ['', '이메일', 'cocoz0707@gmail.com'],
    [f'을 (광고주)', '상호', f'[ {CLIENT["company"]} ]'],
    ['', '대표자', f'[ {CLIENT["representative"]} ]'],
    ['', '사업자등록번호', f'[ {CLIENT["biz_number"]} ]'],
    ['', '주소', f'[ {CLIENT["address"]} ]'],
    ['', '연락처', f'[ {CLIENT["phone"]} ]'],
    ['', '이메일', f'[ {CLIENT["email"]} ]'],
], highlight_cells=hl_cells)
add_text('② 갑과 을은 상호 신뢰를 바탕으로 본 계약을 체결하며, 다음의 조항에 합의한다.')

# === 제2조 ===
add_article('제2조 (계약의 목적)')
add_text('본 계약은 네이버 플레이스 등 지도 서비스 채널 및 검색엔진을 포함한 온라인 마케팅 전반에 걸쳐, 갑이 을에게 제공하는 순위 최적화, 트래픽 관리, 데이터 분석 기반 컨설팅 등 안전보장형 마케팅 서비스에 관한 권리·의무 및 조건을 규정함을 목적으로 한다.')

# === 제3조 ===
add_article('제3조 (서비스의 범위)')
add_text('① 갑이 제공하는 서비스의 범위는 다음 각 호와 같다.')
for s in ['1. 네이버 플레이스 프로필 점검 및 콘텐츠 최적화',
          '2. 검색 노출 최적화 (키워드 전략, 위치 기반 최적화 등)',
          '3. 데이터 분석 기반 트래픽 설계 및 순위 방어·개선',
          '4. 경쟁사 모니터링 및 실시간 대시보드 제공',
          '5. 성과 보고서 제공 (주 1회)',
          '6. 맞춤형 컨설팅 및 전략 조정']:
    add_text('   ' + s)
add_text('② 본 계약에 따른 마케팅 대상 업체(이하 "대상 지점") 및 키워드는 다음과 같다.')

# 지점 테이블 - 전체 하이라이트
branch_hl = set()
for ri in range(len(CLIENT['branches'])):
    for ci in range(4):
        branch_hl.add((ri, ci))
add_table(['번호', '지점명 (플레이스명)', '계약 키워드', '계약 시점 순위'],
          [[f'[ {b[0]} ]', f'[ {b[1]} ]', f'[ {b[2]} ]', f'[ {b[3]} ]'] for b in CLIENT['branches']],
          highlight_cells=branch_hl)

add_text('③ 키워드 변경이 필요한 경우 갑과 을의 서면 합의(이메일, 메신저 등 전자적 기록 포함)를 통해 변경할 수 있다.')
add_text('④ 안전보장형 서비스의 특성상, 갑은 플랫폼 알고리즘 변동에 대응하여 패턴 다양화, 유입 경로 분산, 매체 및 미션 다양화 등의 방식으로 서비스를 수행하며, 구체적 운용 방법은 갑의 전문적 판단에 따른다.')

# === 제4조 - 기간 하이라이트 ===
add_article('제4조 (계약 기간)')

p = doc.add_paragraph()
r = p.add_run('① 본 계약의 기간은 ')
r.font.size = Pt(10)
r = p.add_run(f'[ {CLIENT["period_start"]} ]')
r.font.size = Pt(10)
set_highlight_color(r, LIGHT_YELLOW)
r = p.add_run(' 부터 ')
r.font.size = Pt(10)
r = p.add_run(f'[ {CLIENT["period_end"]} ]')
r.font.size = Pt(10)
set_highlight_color(r, LIGHT_YELLOW)
r = p.add_run(' 까지 ')
r.font.size = Pt(10)
r = p.add_run(f'[ {CLIENT["period_months"]} ]개월')
r.font.size = Pt(10)
set_highlight_color(r, LIGHT_YELLOW)
r = p.add_run('로 한다.')
r.font.size = Pt(10)

add_text('② 계약 만료일 7일 전까지 갑 또는 을이 서면(이메일, 메신저 등 전자적 기록 포함)으로 해지 또는 조건 변경을 통보하지 않는 경우, 동일 조건으로 1개월 단위로 자동 연장된다.')
add_text('③ 자동 연장 시 서비스 이용료는 제5조의 요금이 적용된다.')

# === 제5조 - 가격 하이라이트 ===
add_article('제5조 (대금 및 결제 조건)')
add_text('① 서비스 이용료는 다음과 같다. (부가가치세 10% 별도)')

price_hl = {(0, 1), (0, 2)}
add_table(['구분', '지점당 월 이용료', '4개 지점 합계 (월)'], [
    ['안전보장형 (패키지)', f'[ {CLIENT["price_per_branch"]} ]원', f'[ {CLIENT["price_total"]} ]원'],
], highlight_cells=price_hl)

add_text('② 을은 매월 서비스 개시일 기준 5영업일 이내에 해당 월 서비스 이용료를 갑에게 선결제한다. 단, 첫 달은 계약 체결일로부터 3영업일 이내에 결제한다.')
add_text('③ 결제 방법은 아래 계좌로의 계좌이체를 원칙으로 하며, 갑과 을의 합의로 다른 방법을 사용할 수 있다.')
add_text('    · 은행명: 농협')
add_text('    · 계좌번호: 302-2049-3322-81')
add_text('    · 예금주: 온다마케팅 / 황승현')
add_text('④ 갑은 결제 확인 후 해당 공급일이 속하는 달의 다음 달 10일까지 세금계산서를 발행한다.')
add_text('⑤ 을이 이용료를 10영업일 이상 연체할 경우, 갑은 서면 통보 후 서비스를 일시 중단할 수 있으며, 이 경우 서비스 미제공 기간에 대해 갑은 책임을 지지 아니한다.')
add_text('⑥ 서비스 중단 후 을이 연체금 및 지연이자를 완납하여 서비스를 재개하는 경우, 갑은 별도 재셋업 비용 없이 서비스를 재개한다.')
add_text('⑦ 연체 시 지연이자는 연 12%로 하되, 「상법」 제54조 및 관련 법령에서 정한 상한을 초과하지 아니한다.')

# === 제6조 ===
add_article('제6조 (순위 목표 및 보장 조건)')
add_text('① 안전보장형 서비스는 단기 순위 급상승이 아닌, 플레이스 종합 지수의 안정적 향상을 통한 점진적·지속적 순위 개선을 목표로 한다.')
add_text('② 갑은 각 대상 지점에 대해 내부 지표(N2 스코어 등)의 지속적 상승 추세 유지를 위해 최선을 다한다.')
add_text('③ 다음 각 호의 경우에는 성과 보장의 예외(이하 "예외 사유")로 하며, 갑은 예외 사유 발생 시 지체 없이 을에게 서면 통보하고, 구체적 대응 방안을 함께 제시한다.')
for s in ['1. 네이버 등 플랫폼의 알고리즘·정책 변경으로 인한 순위 변동 (시장 전체적 변동뿐 아니라, 특정 업종·지역에 한정된 변동을 포함한다)',
          '2. 을의 플레이스 계정 정지, 리뷰 삭제 등 플랫폼 제재',
          '3. 을의 영업 중단, 휴업, 폐업 등으로 서비스 수행이 불가능한 경우',
          '4. 천재지변, 전쟁, 정부 규제 등 불가항력 사유',
          '5. 을이 갑의 컨설팅 권고사항을 반복적으로 미이행한 경우']:
    add_text('   ' + s)
add_text('④ 내부 지표의 개선 여부는 갑이 제공하는 자사 분석 솔루션의 N2 스코어를 기준으로 판단하되, 을이 이의를 제기하는 경우 갑과 을이 합의한 제3자 측정 도구(AdLog 등)의 수치를 보조 기준으로 활용할 수 있다.')
add_text('⑤ 갑의 서비스 수행에도 불구하고 2개월 연속 내부 지표의 개선이 전혀 없는 경우, 을은 갑에게 원인 분석 보고서를 요청할 수 있으며, 갑은 7일 이내에 이를 제출하여야 한다. 단, 제③항 각 호의 예외 사유에 해당하는 기간은 위 2개월의 산정에서 제외한다.')
add_text('⑥ 제⑤항의 보고서에도 불구하고 추가 1개월간(예외 사유 기간 제외) 개선이 없는 경우, 을은 위약금 없이 계약을 해지할 수 있다.')

# === 제7조 - 가격 테이블 하이라이트 ===
add_article('제7조 (대상 지점의 추가·삭제 및 가격 조정)')
add_text('① 을이 대상 지점을 추가하고자 하는 경우, 갑과 을은 서면 합의를 통해 추가 지점의 키워드, 이용료를 정한다.')
add_text('② 대상 지점 중 일부가 폐업·휴업·이전 등으로 서비스 대상에서 제외되는 경우, 해당 지점에 대한 서비스는 중단되며, 잔여 지점의 이용료는 다음과 같이 조정한다.')

adj_hl = {(0, 1), (1, 1), (2, 1)}
add_table(['잔여 지점 수', '지점당 월 이용료'], [
    ['3개 지점', f'[ {CLIENT["price_3branch"]} ]원'],
    ['2개 지점', f'[ {CLIENT["price_2branch"]} ]원'],
    ['1개 지점', f'[ {CLIENT["price_1branch"]} ]원'],
], highlight_cells=adj_hl)

add_text('③ 지점 삭제에 따른 이용료 조정은 삭제 통보일이 속하는 달의 다음 달부터 적용한다.')
add_text('④ 잔여 지점이 0개가 되는 경우 본 계약은 자동 종료되며, 제12조의 환불 규정을 준용한다.')

# === 제8조 ===
add_article('제8조 (추가 서비스 및 비용)')
add_text('① 제3조에 명시되지 않은 추가 서비스(블로그 대행, 사진 촬영, 영상 제작, 추가 플랫폼 관리 등) 발생 시, 갑과 을은 별도 서면 합의를 통해 범위와 비용을 확정한다.')
add_text('② 추가 비용은 기존 이용료와 별도로 청구하며, 을의 사전 승인 없이 추가 비용을 발생시킬 수 없다.')

# === 제9조 ===
add_article('제9조 (갑의 의무)')
add_text('① 갑은 본 계약에 따른 마케팅 서비스를 선량한 관리자의 주의의무로 성실히 수행한다.')
add_text('② 갑은 매주 1회 이상 성과 보고서(순위 변동, N2 추이, 경쟁사 비교 등)를 을에게 제공한다.')
add_text('③ 갑은 을의 플레이스 계정 정보, 사업 관련 정보 등 개인정보를 「개인정보 보호법」 및 관련 법령에 따라 안전하게 관리하며, 계약 목적 외 사용·무단 유출·제3자 제공을 하지 아니한다.')
add_text('④ 갑은 플랫폼 정책 변동, 알고리즘 변경 등 서비스에 중대한 영향을 미치는 사항 발견 시 3영업일 이내에 을에게 통보하고 대응 방안을 제시한다.')
add_text('⑤ 갑은 플랫폼의 이용약관 및 관련 법령을 준수하여 서비스를 수행하며, 을의 플레이스에 저품질 패널티 등 불이익이 발생하지 않도록 주의한다.')
add_text('⑥ 갑의 서비스 수행 과정에서 갑의 귀책사유로 을의 플레이스에 저품질 판정, 계정 정지 등 직접적 손해가 발생한 경우, 갑은 해당 월 서비스 이용료를 한도로 손해를 배상한다.')

# === 제10조 ===
add_article('제10조 (을의 의무)')
add_text('① 을은 계약 체결일로부터 7일 이내에 갑의 서비스 수행에 필요한 다음 자료 및 권한을 제공한다.')
for s in ['1. 네이버 플레이스 관리자 접근 권한', '2. 사업자등록증 사본',
          '3. 매장 사진, 메뉴판, 로고 등 마케팅 소재', '4. 기타 갑이 합리적으로 요청하는 자료']:
    add_text('   ' + s)
add_text('② 을은 영업시간 변경, 메뉴 변경, 휴무일, 이벤트 등 플레이스 정보에 영향을 미치는 사항을 사전에 갑에게 통보한다.')
add_text('③ 을은 갑의 컨설팅 권고사항에 성실히 협조한다.')
add_text('④ 을이 제①항의 자료 제공 의무를 계약 체결일로부터 14일 이내에도 이행하지 않는 경우, 갑은 서면 통보 후 서비스 개시를 유보할 수 있으며, 이로 인한 서비스 지연에 대해 갑은 책임을 지지 아니한다.')
add_text('⑤ 을이 제②항 및 제③항의 협조의무에 대해 갑이 2회 이상 서면 독촉 후에도 반복적으로 불이행하는 경우, 갑은 서면 통보로 계약을 해지할 수 있다.')

# === 제11조 ===
add_article('제11조 (계약 해지)')
add_text('① 최소 계약 기간: 본 계약의 최초 1개월은 최소 계약 기간으로 한다. 을이 최초 1개월 내에 해지를 요청하는 경우, 잔여 일수에 대한 환불 없이 해당 월 이용료로 정산한다.')
add_text('② 을의 임의 해지: 을은 최소 계약 기간 경과 후 해지를 원할 경우, 해지 희망일 14일 전까지 갑에게 서면 통보하여야 한다.')
add_text('③ 갑의 해지: 갑은 다음 각 호의 경우 서면 통보로 계약을 해지할 수 있다.')
for s in ['1. 을이 이용료를 30일 이상 연체한 경우',
          '2. 을이 제10조의 협조의무를 갑의 2회 이상 서면 독촉 후에도 반복적으로 불이행한 경우',
          '3. 을의 사업장이 폐업·휴업하여 서비스 수행이 불가능한 경우']:
    add_text('   ' + s)
add_text('④ 갑의 귀책사유로 인한 해지: 갑이 정당한 사유 없이 서비스를 14일 이상 중단하거나 본 계약의 의무를 중대하게 위반한 경우, 을은 즉시 해지 통보 및 잔여 기간에 대한 이용료 전액 환불을 청구할 수 있다.')

# === 제12조 ===
add_article('제12조 (환불)')
add_text('① 을의 임의 해지 시, 갑은 해지일까지의 서비스 이용료를 일할 계산하여 공제한 잔액을 해지일로부터 14영업일 이내에 환불한다. 단, 제11조 ①항에 따라 최소 계약 기간 내 해지 시에는 해당 월 이용료로 정산하며 잔여일수 환불은 적용되지 아니한다.')
add_text('② 갑의 귀책사유로 인한 해지 시, 갑은 미이행 기간에 해당하는 이용료 전액을 7영업일 이내에 환불한다.')
add_text('③ 제10조 ⑤항에 의한 해지(을의 협조의무 불이행) 시, 갑은 진행된 작업일수를 일할 계산하고 계약 잔여금의 20%를 위약금으로 공제한 후 환불한다.')
add_text('④ 환불 시 이미 발행된 세금계산서는 수정세금계산서로 처리한다.')

# === 제13조 ===
add_article('제13조 (손해배상)')
add_text('① 갑 또는 을이 본 계약을 위반하여 상대방에게 손해를 끼친 경우, 귀책 당사자는 상대방에게 직접적이고 통상적인 손해를 배상한다.')
add_text('② 손해배상의 범위는 해당 월 서비스 이용료의 3배를 상한으로 한다. 다만, 고의 또는 중대한 과실로 인한 손해, 개인정보 유출로 인한 손해는 이 상한을 적용하지 아니한다.')
add_text('③ 제②항 단서의 개인정보 유출에 관하여, 갑이 본 계약에 따라 보유·처리하는 을의 개인정보의 범위는 제10조 ①항 각 호의 자료에 한정되며, 갑의 손해배상 책임은 해당 범위 내의 정보에 대해서만 적용된다.')
add_text('④ 갑은 플랫폼(네이버 등)의 정책 변경, 알고리즘 변동, 서버 장애 등 갑이 통제할 수 없는 외부 요인으로 인한 순위 변동·서비스 중단에 대해서는 책임을 지지 아니한다.')

# === 제14조 ===
add_article('제14조 (지식재산권 및 콘텐츠 권리)')
add_text('① 갑이 본 계약의 이행 과정에서 을의 요청에 따라 제작한 광고 콘텐츠(이미지, 영상, 문구, 디자인 등)의 저작재산권은 을에게 귀속된다.')
add_text('② 갑은 제①항의 콘텐츠를 자사 포트폴리오·홍보 목적으로 비독점적으로 사용할 수 있다. 다만, 을의 브랜드 이미지에 중대한 영향을 미칠 수 있는 사용은 을의 사전 동의를 받는다.')
add_text('③ 을이 특정 콘텐츠의 2차 활용을 제한하고자 할 경우, 별도 서면 합의로 정한다.')
add_text('④ 갑은 본 계약에서 제작된 콘텐츠를 제3자에게 양도·판매할 수 없다.')
add_text('⑤ 계약 해지 후에도 갑이 기 제작·게시한 콘텐츠는 을의 별도 삭제 요청이 없는 한 유지된다.')

# === 제15조 ===
add_article('제15조 (비밀유지)')
add_text('① 갑과 을은 본 계약의 체결 및 이행 과정에서 알게 된 상대방의 영업비밀, 고객정보, 마케팅 전략, 기술 정보 등 일체의 비밀정보를 계약 기간 중은 물론 계약 종료 후 2년간 제3자에게 누설·공개·제공하지 아니한다.')
add_text('② 다음 각 호에 해당하는 정보는 비밀정보에서 제외한다.')
for s in ['1. 공지 시 이미 공개적으로 알려진 정보',
          '2. 수령 당사자가 비밀유지 의무 없이 적법하게 보유하고 있던 정보',
          '3. 법령 또는 관할 법원·행정기관의 명령에 의해 공개가 요구되는 정보']:
    add_text('   ' + s)
add_text('③ 갑은 을의 사업장 정보, 매출 데이터, 고객 데이터 등을 「개인정보 보호법」에 따라 처리하며, 계약 종료 시 을의 요청에 따라 30일 이내에 파기 또는 반환한다.')

# === 제16조 ===
add_article('제16조 (계약의 양도 제한)')
add_text('① 갑과 을은 상대방의 사전 서면 동의 없이 본 계약상의 권리·의무를 제3자에게 양도·이전하거나 담보로 제공할 수 없다.')

# === 제17조 ===
add_article('제17조 (통지)')
add_text('① 본 계약에 따른 모든 통지는 제1조에 기재된 주소, 이메일 또는 상호 합의한 메신저(카카오톡, 문자 등)를 통해 발송하며, 다음 각 호의 시점에 도달한 것으로 본다.')
add_text('   1. 이메일·메신저: 발송 시점')
add_text('   2. 등기우편: 발송일로부터 3영업일')
add_text('② 연락처 변경 시 상대방에게 7일 이내에 통보하여야 하며, 통보하지 않음으로 인한 불이익은 변경 당사자가 부담한다.')

# === 제18조 ===
add_article('제18조 (불가항력)')
add_text('① 천재지변, 전쟁, 내란, 법령 변경, 정부 규제, 플랫폼의 서비스 중단·정책 변경, 해킹, 대규모 네트워크 장애 등 당사자의 합리적 통제 범위를 벗어난 사유(이하 "불가항력")로 인하여 본 계약의 이행이 지연·불가능한 경우, 해당 당사자는 그 범위 내에서 책임을 지지 아니한다.')
add_text('② 불가항력 사유 발생 시 해당 당사자는 지체 없이 상대방에게 그 사실과 예상 기간을 통보하여야 한다.')
add_text('③ 불가항력 사유가 60일 이상 지속되는 경우, 어느 당사자든 위약금 없이 계약을 해지할 수 있다.')

# === 제19조 ===
add_article('제19조 (분쟁 해결)')
add_text('① 본 계약과 관련하여 분쟁이 발생한 경우, 갑과 을은 우선 성실한 협의를 통해 해결한다.')
add_text('② 협의로 해결되지 않는 경우, 갑의 본사 소재지를 관할하는 법원을 제1심 전속 관할법원으로 한다.')
add_text('③ 본 계약의 해석 및 적용에 관하여는 대한민국 법률을 준거법으로 한다.')

# === 제20조 ===
add_article('제20조 (계약의 완전합의)')
add_text('① 본 계약서는 본 건에 관한 갑과 을 사이의 완전한 합의를 구성하며, 본 계약 체결 이전의 구두·서면 합의, 제안서, 양해각서 등은 본 계약에 저촉되는 범위 내에서 효력을 상실한다.')
add_text('② 본 계약의 변경·수정은 갑과 을의 서면 합의에 의해서만 효력이 발생한다.')

# === 제21조 서명 - 날짜+을 정보 하이라이트 ===
add_article('제21조 (계약 서명)')
add_text('본 계약서의 내용에 대해 갑과 을 당사자는 충분히 협의하고, 모든 조항을 이해·동의한다.')
add_text('이에 아래 서명·날인을 함으로써 본 계약의 효력이 발생한다.')
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('계약 체결일: ')
r.font.size = Pt(10)
r.font.bold = True
r = p.add_run(f'[ {CLIENT["contract_date_text"]} ]')
r.font.size = Pt(10)
r.font.bold = True
set_highlight_color(r, LIGHT_YELLOW)

doc.add_paragraph()

# 갑 서명란
add_text('갑 (마케팅 제공업체)')
add_table(['항목', '내용'], [
    ['상호', '온다마케팅'],
    ['대표자', '황승현'],
    ['사업자등록번호', OUR_BIZ_NUMBER],
    ['주소', '서울시 강남구 테헤란로 82길 15, 3층 304호'],
])
add_text('서명/날인: ____________________')
doc.add_paragraph()

# 을 서명란 - 전체 하이라이트
add_text('을 (광고주)')
sign_hl = {(0, 1), (1, 1), (2, 1), (3, 1)}
add_table(['항목', '내용'], [
    ['상호', f'[ {CLIENT["company"]} ]'],
    ['대표자', f'[ {CLIENT["representative"]} ]'],
    ['사업자등록번호', f'[ {CLIENT["biz_number"]} ]'],
    ['주소', f'[ {CLIENT["address"]} ]'],
], highlight_cells=sign_hl)
add_text('서명/날인: ____________________')
doc.add_paragraph()
add_text('본 계약서는 2부 작성하여 갑과 을이 각 1부씩 보관한다.')

# 저장
outpath = 'docs/contracts/온라인_마케팅_계약서_안전보장형_v6.docx'
doc.save(outpath)
print(f'✅ v6 계약서 생성 완료 (하이라이트 적용): {outpath}')